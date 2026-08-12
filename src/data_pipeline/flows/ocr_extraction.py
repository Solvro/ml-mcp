import logging
import os
import re
from pathlib import Path

import pymupdf
import pytesseract
from docx import Document
from PIL import Image
from prefect import get_run_logger, task

MIN_EXTRACTED_TEXT_CHARS = 50
DEFAULT_PDF_RENDER_SCALE = 2.0


def _get_logger():
    try:
        return get_run_logger()
    except Exception:
        return logging.getLogger(__name__)


def _normalize_text(text: str) -> str:
    collapsed = re.sub(r"[ \t]+", " ", text or "")
    stripped = "\n".join(line.strip() for line in collapsed.split("\n"))
    return re.sub(r"\n{3,}", "\n\n", stripped).strip()


def _is_text_sufficient(text: str, min_chars: int) -> bool:
    return len(text.strip()) >= min_chars


def _env_number(name: str, default: int | float) -> int | float:
    """Read a positive number from env; warn and fall back on invalid input."""
    raw = os.getenv(name, "").strip()
    if not raw:
        return default
    try:
        parsed = type(default)(raw)
    except ValueError:
        _get_logger().warning("Invalid %s=%r; using default %s", name, raw, default)
        return default
    if parsed <= 0:
        _get_logger().warning("Non-positive %s=%r; using default %s", name, raw, default)
        return default
    return parsed


def _ocr_pdf_page(page: pymupdf.Page, scale: float, lang: str) -> str:
    """Render one page to an image and run Tesseract on it."""
    pixmap = page.get_pixmap(matrix=pymupdf.Matrix(scale, scale), alpha=False)
    image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
    return pytesseract.image_to_string(image, lang=lang)


def _extract_page_text(
    page: pymupdf.Page,
    page_index: int,
    min_chars: int,
    scale: float,
    lang: str,
) -> str | None:
    """Return page text, or None when extraction failed and the page must be retried."""
    logger = _get_logger()

    try:
        native_text = _normalize_text(page.get_text())
    except Exception as exc:
        logger.warning("Native text extraction failed for page %d: %s", page_index + 1, exc)
        native_text = ""

    if _is_text_sufficient(native_text, min_chars):
        return native_text

    try:
        ocr_text = _normalize_text(_ocr_pdf_page(page, scale, lang))
    except Exception as exc:
        logger.warning("OCR failed for page %d, skipping: %s", page_index + 1, exc)
        return None

    return ocr_text or native_text


def _extract_pdf_pages(source_id: str, file_path: Path) -> list[tuple[str, str]]:
    min_chars = _env_number("OCR_MIN_TEXT_CHARS", MIN_EXTRACTED_TEXT_CHARS)
    scale = _env_number("OCR_PDF_RENDER_SCALE", DEFAULT_PDF_RENDER_SCALE)
    lang = os.getenv("OCR_LANG", "pol+eng").strip() or "pol+eng"

    output: list[tuple[str, str]] = []

    with pymupdf.open(str(file_path)) as doc:
        for page_index, page in enumerate(doc):
            final_text = _extract_page_text(page, page_index, min_chars, scale, lang)
            if not final_text:
                continue
            output.append((f"{source_id}#page={page_index + 1}", final_text))

    return output


def _extract_docx_text(file_path: Path) -> str:
    document = Document(str(file_path))
    parts: list[str] = []

    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_values = [
                cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()
            ]
            if row_values:
                parts.append(" | ".join(row_values))

    return _normalize_text("\n".join(parts))


@task
def ocr_extraction(acquired_documents: list[dict[str, str]]) -> list[tuple[str, str]]:
    """Extract text from acquired documents, using native text and OCR as needed."""
    logger = _get_logger()

    if not acquired_documents:
        return []

    tesseract_cmd = os.getenv("TESSERACT_CMD", "").strip()
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    extracted: list[tuple[str, str]] = []

    for item in acquired_documents:
        source_id = str(item.get("source_id", "")).strip()
        path_raw = str(item.get("path", "")).strip()

        if not source_id or not path_raw:
            logger.warning("Skipping invalid acquisition item: %s", item)
            continue

        file_path = Path(path_raw)
        if not file_path.exists() or not file_path.is_file():
            logger.warning("Skipping missing file for source_id=%s path=%s", source_id, file_path)
            continue

        suffix = file_path.suffix.lower()

        try:
            if suffix == ".pdf":
                extracted.extend(_extract_pdf_pages(source_id, file_path))
            elif suffix in {".txt", ".md"}:
                text = _normalize_text(file_path.read_text(encoding="utf-8", errors="ignore"))
                if text:
                    extracted.append((source_id, text))
            elif suffix == ".docx":
                text = _extract_docx_text(file_path)
                if text:
                    extracted.append((source_id, text))
            else:
                logger.debug("Skipping unsupported extension during extraction: %s", file_path)
        except Exception as exc:
            logger.warning("Extraction failed for source_id=%s (%s): %s", source_id, file_path, exc)

    return extracted
